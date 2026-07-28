"""Word version of the brief, for anyone who needs to edit or comment on it.

    python docs/build-brief-docx.py

Same content and same discipline as the PDF: measured, modelled and assumed are
labelled where the figure appears, not in a footnote nobody reads.
"""
import os
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor, Inches

OUT = os.path.join(os.path.dirname(__file__), "Warrant-Brief.docx")

INK = RGBColor(0x12, 0x21, 0x1F)
SOFT = RGBColor(0x3F, 0x56, 0x54)
FAINT = RGBColor(0x7A, 0x90, 0x8E)
TEAL = RGBColor(0x2C, 0x7A, 0x87)

doc = Document()
for section in doc.sections:
    section.left_margin = section.right_margin = Inches(0.9)
    section.top_margin = section.bottom_margin = Inches(0.85)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(10)
normal.font.color.rgb = SOFT
normal.paragraph_format.space_after = Pt(7)


def h(text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = INK if level == 1 else TEAL
        run.font.name = "Georgia" if level == 1 else "Calibri"
        run.font.size = Pt(19 if level == 1 else 12)
    return p


def para(text, bold_prefix=None, italic=False, size=10, color=SOFT):
    p = doc.add_paragraph()
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
        r.font.color.rgb = INK
        r.font.size = Pt(size)
    r = p.add_run(text)
    r.italic = italic
    r.font.size = Pt(size)
    r.font.color.rgb = color
    return p


def bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
        r.font.color.rgb = INK
    r = p.add_run(text)
    r.font.color.rgb = SOFT
    return p


def table(rows, widths=None, header=True):
    t = doc.add_table(rows=0, cols=len(rows[0]))
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, row in enumerate(rows):
        cells = t.add_row().cells
        for j, value in enumerate(row):
            cells[j].text = ""
            p = cells[j].paragraphs[0]
            p.paragraph_format.space_after = Pt(2)
            r = p.add_run(str(value))
            r.font.size = Pt(8.5)
            r.bold = header and i == 0
            r.font.color.rgb = INK if (header and i == 0) else SOFT
    if widths:
        for j, w in enumerate(widths):
            for row in t.rows:
                row.cells[j].width = Inches(w)
    doc.add_paragraph()
    return t


def note(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.italic = True
    r.font.size = Pt(8.5)
    r.font.color.rgb = FAINT
    return p


# ------------------------------------------------------------------- cover
title = doc.add_paragraph()
tr = title.add_run("Warrant")
tr.font.name = "Georgia"
tr.font.size = Pt(38)
tr.font.color.rgb = INK
tr.bold = True

sub = doc.add_paragraph()
sr = sub.add_run("Verified-fact estate settlement, built for Alix")
sr.font.name = "Georgia"
sr.font.size = Pt(15)
sr.font.color.rgb = TEAL

para("Every fact carries the sentence that justifies it. Verification is deterministic — "
     "no model is ever asked whether a model told the truth.", size=11)

table([
    ["Team", "Mitansh and Egor"],
    ["Built for", "Alix — Agents of Administration, Track 3 (with Track 1 intact)"],
    ["Repository", "github.com/mit37/warrant"],
    ["Date", "28 July 2026"],
    ["Status", "463 tests. Anvil live. Extraction run against a real model."],
], widths=[1.4, 5.0], header=False)

doc.add_page_break()

# ----------------------------------------------------------------- summary
h("What this is, in one page")
para("Warrant is a decision engine for estate settlement in which no claim reaches a filed "
     "document without evidence behind it. A model may propose any fact it likes; it cannot get "
     "that fact into a decision without producing a verbatim quotation that a deterministic "
     "checker locates in a document we actually hold. Facts that fail are quarantined — visible "
     "to a reviewer, unreachable by the rules engine.", size=11)

table([["463 tests", "4 state packs", "58 CA counties", "171 field mappings", "41/41 verified"]],
      header=False)

h("The four invariants", 2)
bullet("A quotation the checker cannot locate is quarantined.", "No fact without a warrant. ")
bullet("Models extract. Rules decide. The evaluator is pure.", "No model in the decision path. ")
bullet("Totals are computed in code under a cited statutory rule.", "No model arithmetic. ")
bullet("A predicate over a fact we do not hold is unknown, never false. A blocked rule names "
       "what it needs, so gap detection is free rather than a feature.", "Three-valued logic. ")

h("What it does that nothing else here does", 2)
bullet("reads the words printed around each field geometrically, has a model propose the "
       "mapping, and refuses any mapping whose claimed label is not on the page.",
       "Onboards a government form it has never seen — ")
bullet("infers a life policy from a $14.32 monthly debit that appears in no document.",
       "Finds assets nobody mentioned — ")
bullet("every citation carries a re-check horizon, and where a statute publishes its amendment "
       "cycle, the exact date.", "Knows when its own answers expire — ")
bullet("of twenty form/estate pairs, six are correctly withheld and one is blocked rather than "
       "answered.", "Declines, with a reason — ")

doc.add_page_break()

# ----------------------------------------------------------- functionality
h("Functionality, in full")

h("Extraction and verification", 2)
table([
    ["Capability", "What it does", "State"],
    ["Live extraction", "Documents to candidate facts, each with the sentence it rests on",
     "Run: 9 docs, 41 facts, 41 verified, $0.08"],
    ["Deterministic verifier", "Normalised exact match, then n-gram coverage. No model involved",
     "Inherited, unchanged, load-bearing"],
    ["Numeric guard", "A quote whose figures are not in the document cannot verify",
     "Added after a fabricated fee passed"],
    ["Quarantine", "Failed facts are visible and structurally unreachable by rules",
     "Enforced by type"],
    ["Record warrants", "Imported facts cite system, record and field path", "Checked on admission"],
], widths=[1.3, 2.9, 2.2])

h("Jurisdiction", 2)
table([
    ["Capability", "Detail"],
    ["Rule packs", "California, New York, Pennsylvania, Texas — each sourced to the state's own legislature"],
    ["County pack", "All 58 California counties, with fee variation (San Francisco is $450, not $435)"],
    ["Honest gaps", "Each pack carries an explicit list of what could NOT be sourced"],
    ["No fallback", "An uncovered state gets 'no rule pack' and a hold, never another state's answers"],
    ["Freshness", "Every citation has a re-check horizon; scheduled statutory changes beat observed stability"],
    ["Compiler", "New counties compile from their own published text, model-proposed and quote-verified"],
], widths=[1.3, 5.1])

h("Forms — Track 3", 2)
table([
    ["Capability", "Detail"],
    ["Geometry extraction", "Widget rectangles plus the words printed around each, read off the page"],
    ["Discovered mapping", "Model proposes; the same matcher that checks facts checks the cited label"],
    ["Refusal", "Two proposals cited labels not printed where they claimed. Recorded, not deleted"],
    ["Adjudications", "Human corrections live in an override layer, re-applied and re-verified"],
    ["Applicability", "Which forms apply is a rule with a citation, not a hardcode"],
    ["Local fill", "Real filled PDFs, every field read back"],
    ["Anvil", "Live. 139/139 bindings joined on three forms; values confirmed on the page"],
], widths=[1.3, 5.1])

h("Discovery, dispatch and audit", 2)
table([
    ["Capability", "Detail"],
    ["Recurring detection", "157 transactions to 12 recurring charges, across descriptor drift"],
    ["Asset inference", "Three assets found in no document: life policy, property policy, safe deposit box"],
    ["Suppression", "An insurer already in the ledger is deliberately not re-reported"],
    ["Hold detection", "Classifies the line from audio; transcription starts only when a person speaks"],
    ["Bounded speech", "The agent may state only verified facts within the prepared scope"],
    ["Dependency replay", "Supersede one fact; only decisions that read it are re-evaluated"],
    ["Approval gate", "Keyed on reversibility and blast radius, not a blanket confidence score"],
], widths=[1.3, 5.1])

doc.add_page_break()

# -------------------------------------------------------------- evidence
h("What has actually run")
para("Measured means a command in the repository produced it. Modelled means it was calculated "
     "from assumptions that are named. Nothing here is an estimate presented as an observation.",
     size=11)

table([
    ["Result", "Figure", "Evidence"],
    ["Live extraction", "41 proposed, 41 verified, 0 quarantined", "measured, replayable"],
    ["Extraction cost", "$0.08 for nine documents", "measured"],
    ["Field mappings", "171 verified across 4 forms, 3 refused", "measured"],
    ["Anvil chain", "139/139 bindings, values read back", "measured, live API"],
    ["Filled PDFs", "14 documents, 270 fields", "measured"],
    ["Forms correctly withheld", "6 of 20 pairs, each with a reason", "measured"],
    ["Model agreement", "87 of 94; 3 of 7 disagreements were real errors", "measured"],
    ["Recurring bleed found", "$11,963/yr still leaving the estate", "measured"],
], widths=[1.6, 2.4, 2.4])

h("The result we are least comfortable with, stated anyway", 2)
para("Two mappings passed verification and were still wrong. One put the signer's name in "
     "“Title, if applicable”; the other answered “are the assets in the custody of the court?” "
     "from a field meaning “does a court case exist”. Both cited labels genuinely printed on the "
     "form. Locating a quote proves the model did not invent its evidence; it does not prove the "
     "reasoning is right. We found these by rendering the filled PDF and reading it, and they "
     "are recorded as adjudications rather than quietly fixed.")

doc.add_page_break()

# ------------------------------------------------------------------ money
h("Financial model")
para("The widely-quoted 570–900 hours is what the family spends. Alix's P&L turns on what its "
     "specialist spends. Reducing the first is marketing; reducing the second is margin. This "
     "model addresses the second and says so.", size=11)

h("Inputs", 2)
table([
    ["Input", "Value", "Status"],
    ["Fully-loaded specialist hour", "$70",
     "Alix's own figure — salary, benefits, payroll tax. Corroborated: BLS OEWS wages against "
     "the ECEC benefit load give $58 low, $74 central, $102 high"],
    ["Alix revenue per estate", "$9,000 min", "Alix published pricing — 1% of assets, $9,000 floor"],
    ["Probate referee commission (CA)", "0.1%", "Cal. Prob. Code 8961(a); min $75, max $10,000"],
    ["Hours to onboard one form today", "4 h",
     "ASSUMPTION. Nobody has been timed — the single figure most worth replacing"],
], widths=[1.7, 0.9, 3.8])

h("Per-estate saving, at $70/hr", 2)
table([
    ["Lever", "Basis", "Hours", "Value"],
    ["Form filling — 270 fields auto-populated", "measured", "2 – 4", "$140 – 280"],
    ["Hold time — 34 min per call, ~6 calls", "modelled", "3.4", "$238"],
    ["Jurisdiction lookup instead of research", "modelled", "2 – 5", "$140 – 350"],
    ["Recurring-charge shutdown", "measured", "1 – 2", "$70 – 140"],
    ["Appraisal — cash self-appraised under 8901", "measured", "—", "$460 on the LA sample"],
    ["Total per estate", "", "8 – 14", "$560 – 980"],
], widths=[2.6, 0.9, 0.8, 2.1])
para("Against $9,000 of revenue that is 6 to 11 points of gross margin, on a business whose "
     "alternative is hiring. And it compounds: the same rule pack and field map serve every "
     "subsequent estate in that jurisdiction at no additional cost, so the second estate in a "
     "county is cheaper than the first and the hundredth is nearly free.")

h("The larger number: form onboarding", 2)
table([
    ["Scenario", "New forms", "Specialist time today", "Cost at $70", "With Warrant"],
    ["Pilot", "20", "80 h", "$5,600", "model cost + review"],
    ["Year one", "200", "800 h", "$56,000", "model cost + review"],
    ["At scale", "1,000", "4,000 h", "$280,000", "model cost + review"],
], widths=[1.0, 0.7, 1.4, 0.9, 2.4])
note("Four hours per form is an assumption, not a measurement. Ian is the right person to "
     "replace it, and one sentence from him makes this table defensible in a way our own "
     "benchmark never will be.")

h("Three-year shape", 2)
table([
    ["", "Year 1", "Year 2", "Year 3"],
    ["Estates processed", "250", "1,200", "4,000"],
    ["Per-estate saving at $770 midpoint", "$192,500", "$924,000", "$3,080,000"],
    ["New forms onboarded that year", "200", "400", "900"],
    ["Onboarding saving at $70/hr", "$56,000", "$112,000", "$252,000"],
    ["Total", "$248,500", "$1,036,000", "$3,332,000"],
    ["Jurisdictions covered", "4 states", "15 states", "40 states"],
], widths=[2.2, 1.4, 1.4, 1.4])
note("Estate volumes are illustrative. Alix does not publish how many estates it settles and we "
     "will not invent the figure — replace the top row and every row below follows "
     "arithmetically. The saving rate and the hourly cost are the parts we stand behind.")

doc.add_page_break()

# ------------------------------------------------------------- assessment
h("Honest assessment")

h("What is genuinely strong", 2)
bullet("It rejects. A fabricated policy is quarantined; a substituted fee no longer passes; "
       "three form mappings were refused.", "The verification gate is real and load-bearing. ")
bullet("No human placed a field. 258 widgets across four forms, and the marginal cost of the "
       "next form is a model call and a review.", "Form onboarding scales the way Alix needs. ")
bullet("Compiled, cited, versioned, diffable, zero tokens at decision time — and it declines "
       "rather than guessing outside its coverage.", "The jurisdiction model is the right shape. ")
bullet("Inference from payment traces, not extraction. The industry currently tells families to "
       "do this by hand.", "Asset discovery is genuinely novel. ")

h("What is weak, and by how much", 2)
table([
    ["Weakness", "Severity", "Detail"],
    ["Four states of fifty", "High", "Architecture scales; content does not yet. Research-bound"],
    ["Three courts researched", "High", "Lazy acquisition helps, but it is still a queue"],
    ["Nobody has used it", "High", "No specialist has touched it. No user testing"],
    ["Verification has a ceiling", "Medium", "Catches fabricated evidence, not wrong reasoning"],
    ["Voice never placed a call", "Medium", "Designed and tested, but unproven in the field"],
    ["One estate, nine documents", "Medium", "Extraction proven narrowly. No OCR'd scans"],
    ["DL 142 does not join Anvil", "Low", "Anvil's detection disagrees with that PDF"],
    ["Economics are modelled", "Medium", "Two of five saving levers are calculated, not observed"],
], widths=[1.7, 0.8, 3.9])

h("Implementation path", 2)
table([
    ["Phase", "Scope", "Effort", "Integration risk"],
    ["1. Shadow run", "Point the importer at real estates", "1–2 weeks", "None — reads their format"],
    ["2. Form onboarding", "They hand over PDFs, we return maps and Casts", "2–4 weeks", "Low — their Anvil account"],
    ["3. Jurisdiction service", "Rule packs behind an API", "4–8 weeks", "Medium — a new dependency"],
    ["4. Dispatch", "Hold detection, mid-call release", "8+ weeks", "High — needs legal sign-off"],
], widths=[1.3, 2.4, 0.9, 1.8])
para("Phase 2 pays for itself and touches nothing. It is the highest return and the lowest risk, "
     "it addresses the deficiency Alix themselves named, and it needs no changes on their side "
     "beyond handing over a folder of PDFs.")

h("How it compares", 2)
table([
    ["", "What they do well", "What is missing", "Our position"],
    ["Anvil", "Excellent PDF fill and e-signature", "Fields placed by a human, per form",
     "Complementary — we generate and verify the mapping"],
    ["EstateExec, Trust & Will", "Consumer checklists", "No jurisdiction engine", "Different user"],
    ["Generic LLM + RAG", "Fast to stand up", "No warrant, no effective dates, similarity where you need lookup",
     "The approach we deliberately rejected"],
    ["Alix today", "Real expertise and clients", "Volume and defensibility, by their own account",
     "Infrastructure underneath, not a replacement"],
], widths=[1.1, 1.5, 1.9, 1.9])

doc.add_page_break()

# ----------------------------------------------------------------- close
h("What we need")
table([
    ["Item", "Why", "Blocking?"],
    ["Twilio credentials", "The voice agent has never placed a call", "Yes, for dispatch"],
    ["SMTP / fax credentials", "Mid-call document release", "Yes, for dispatch"],
    ["Ian: hours to onboard one form", "Replaces the assumption the onboarding table rests on",
     "No, but highest-value"],
    ["Priority counties", "Lazy acquisition builds only what the book touches", "No"],
    ["Alix estate volume", "Turns the three-year table from illustrative into real", "No"],
], widths=[1.6, 3.0, 1.8])

h("Two things we would say out loud", 2)
para("It traces to the founder's own estate, was later restated as a population average, and is "
     "now cited back by journalists as an industry estimate. It appears on at least six Alix "
     "pages with no citation. The nearest real research says 570, from a survey with a stated "
     "sample and error margin. We would use Alix's published range verbatim — 600 to 900 — "
     "because a range cannot be attacked as a false point estimate, and its low end brackets the "
     "survey.", bold_prefix="The 900-hour figure does not survive its source. ")
para("It makes the system accountable, which is a smaller claim and a more defensible one. "
     "Everything this product does is designed so that when it is wrong, somebody can see "
     "exactly where and why.",
     bold_prefix="We would not claim the verifier makes the system correct. ")

q = doc.add_paragraph()
qr = q.add_run("“Tomorrow, if your system told me that I have to go through formal probate, you "
               "need to track me every decision that was made. You need to tell me where you got "
               "your legal information.”")
qr.italic = True
qr.font.size = Pt(11)
qr.font.color.rgb = INK
q.paragraph_format.left_indent = Inches(0.3)
q.alignment = WD_ALIGN_PARAGRAPH.LEFT

note("— Soren, Rules as Data track")
para("That is the whole specification, and it is the one we built to.")

doc.save(OUT)
print(f"wrote {OUT}  ({os.path.getsize(OUT) // 1024} KB)")
